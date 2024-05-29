import docx

def create_question_doc(questions, doc_path):
  """Creates a Word document with formatted questions and answer choices.

  Args:
      questions: A list of strings, where each string represents a question
                  or answer choice.
      doc_path: The path to save the Word document.
  """

  document = docx.Document()

  # Add title with heading style
  document.add_heading('Questions', level=1)

  # Add each question and answer choice as a paragraph, handling indentation
  for line in questions:
    paragraph = document.add_paragraph()
    paragraph_text = line.strip()  # Remove leading/trailing whitespace

    # Apply formatting based on content structure
    if paragraph_text.isdigit() or paragraph_text.startswith('Match'):
      paragraph.style = document.styles['Heading 2']  # Level 2 heading for titles
    elif paragraph_text.startswith(' ') * 4:  # Multiple spaces for indentation
      paragraph.style = document.styles['List Bullet 2']  # Level 2 bullet point
    else:
      paragraph.style = document.styles['Normal']  # Normal style

    paragraph.text = paragraph_text

  # Save the document
  document.save(doc_path)

 # Sample questions list with indentation and formatting hints
questions = [
  "Which one is the color of the sky?",
  "  a) Red",
  "  b) Blue",
  "  c) Yellow",
  "",
  "2. Any instrument that is held in the hand to do a particular work is a …………",
  "  a) tool",
  "  b) material",
  "  c) table",
  "",
 "",
    "3. Which one is a fruit?",
    "   a) onion",
    "   b) mango",
    "   c) ball",
    "",
    "4. Which animal can fly?",
    "   a) mouse",
    "   b) bird",
    "   c) elephant",
    "",
    "5. What do we draw on?",
    "   a) paper",
    "   b) car",
    "   c) desk",
    "",
    "6. What do we use to colour?",
    "   a) crayon",
    "   b) pen",
    "   c) stone",
    "",
    "7. …………… things are things created by God.",
    "   a) artificial",
    "   b) man made",
    "   c) natural",
    "",
    "8. List two (2) examples of tools.",
    "   a. ",
    "   b. ",
    "",
    "9. Which of the following is a natural thing",
    "   a) rubber bowl",
    "   b) book",
    "   c) stone",
    "",
    "10. Circle the odd one",
    "   a) Prosper",
    "   b) pencil",
    "   c) colour",
    "",
    "Match the following images with their corresponding names:",
    "   Fabric",
    "   Painting Brush",
    "   Rope",
    "   Wooden hammer",
    "   Scissors"
]

# Set the document path (replace with your desired location)
doc_path = r'C:\Users\orbia\OneDrive\Desktop\Dede\Questions.docx'  # Adjust path as needed

create_question_doc(questions, doc_path)

